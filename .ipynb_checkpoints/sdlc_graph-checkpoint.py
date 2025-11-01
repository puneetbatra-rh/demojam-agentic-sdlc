import os
from dotenv import load_dotenv
import re
load_dotenv() ## aloading all the environment variable

AUTH_TOKEN=os.getenv("LLAMA_API_KEY")
MODEL=os.getenv("LLAMA_MODEL")
INFERENCE_URL_BASE=os.getenv("LLAMA_URL")


# building Graph
from langgraph.graph import END, StateGraph, START
from langchain_core.prompts import PromptTemplate
from pydantic import BaseModel, Field
from typing import Literal
from langchain_core.output_parsers import StrOutputParser
from langgraph.checkpoint.memory import MemorySaver
import uuid
from pprint import pprint
from docx import Document

from langchain_openai import ChatOpenAI
import os

# The 'model_name' here is often a placeholder or the deployment name
# and is required by the ChatOpenAI class.
llm = ChatOpenAI(
    model=MODEL,
    openai_api_base=INFERENCE_URL_BASE,
    openai_api_key=AUTH_TOKEN
    # Disable SSL verification if you are using a self-signed or internal certificate
    # client_kwargs={"verify": False}
)


# Data modeling
# Define Graph State
from typing import List, Dict
from typing_extensions import TypedDict

class State(TypedDict):
    """
    Represents in the stage of our graph.

    Attributes:
        requirements: user requirement
        user_stories: user stories
        design_document: design document functional and technical
        code: code
        test_cases: test cases
        qa_testing: qa testing
        monitoring: monitoring
        feedback: feedback
        status: status
    """
    requirements: str
    user_stories: List[str]
    user_story_status: Literal["Approve", "Denied"]
    user_story_feedback: List[str]
    design_document_review_status: Literal["Approve", "Denied"]
    design_document_review_feedback: List[str]
    design_document: Dict[List[str], List[str]]
    code: str
    code_review_status: str
    code_review_feedback: List[str]
    security_review_status: str
    security_review_feedback: str
    test_cases: str
    test_cases_review_status: str
    test_cases_review_feedback: List[str]
    qa_review_status: str
    qa_review_feedback: List[str]
    deployment: str

class UserStories(BaseModel):
    stories: List[str]


class DesignDocument(BaseModel):
    functional:List[str] = Field(description="Functional Design Document")
    technical:List[str] = Field(description="Technical Design Document")
    assumptions: List[str] = Field(default=[], description="Design assumptions and constraints")
    open_questions: List[str] = Field(default=[], description="Open questions and risks")

class Review(BaseModel):
    review: str = Field(
        description="Detailed feedback that provides specific, actionable insights about strengths and weaknesses. Should include concrete suggestions for improvement with clear reasoning behind each point. For code reviews, include comments on quality, readability, performance, and adherence to best practices. For design documents, address completeness, clarity, and technical feasibility."
    )
    status: Literal["Approve", "Denied"]

class GenerateCode(BaseModel):
    generated_code: str = Field(
        description="Generated code in the format mentioned in the prompt."
    )

class TestCases(BaseModel):
    cases: str


def user_input_requirements(state: State):
    return state


def save_user_stories_to_txt(state: dict, output_dir="artifacts", filename="user_stories.txt"):
    """
    Saves the list of user stories from the state into a formatted text file in the artifacts directory.

    Args:
        state (dict): The LangGraph state containing the 'user_stories' key.
        output_dir (str): The directory where the file should be saved.
        filename (str): The name of the output text file.
    """
    # Ensure the directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Get user stories
    user_stories = state.get("user_stories", [])

    if not user_stories:
        print("⚠️ No user stories found in state.")
        return

    # Format and write to file
    output_path = os.path.join(output_dir, filename)
    with open(output_path, "w") as f:
        f.write("📘 User Stories\n")
        f.write("="*40 + "\n\n")
        for i, story in enumerate(user_stories, 1):
            f.write(f"{i}. {story}\n")

    print(f"✅ User stories saved to: {output_path}")


def auto_generate_user_stories(state:State):
    # feedback_points = state.get('user_story_feedback', "")

    if state["requirements"] == "":
        return {"error": "Please enter requirement before generating user stories!!"}

    prompt_user_stories = PromptTemplate(

        template = """
        You are a seasoned Agile Product Manager with deep expertise in crafting user-centric user stories.

        Your task is to generate 6 well-defined user stories based on the following product requirement:
        "{requirements}".

        If feeback points are present then also consider feedback points also which generating the uesr stories.
        "{feedback_points}"

        Guidelines for each user story:
        - Use this format: *As a <type of user>, I want to <goal or feature> so that <benefit or reason>.*
        - Focus on the user's perspective and their value—not technical implementation details.
        - Ensure each story is:
        - Concise but informative
        - Independent from the others
        - Testable, with implied or clear acceptance criteria
        - Representing a different functional aspect of the requirement

        Think broadly about user roles (e.g., admin, end user, guest, etc.) to provide well-rounded coverage of the application's core functionality.

        Return only the 6 user stories in a numbered list, with no additional explanation.
        """,
        input_variables=["requirements"]
    )
    
    chain_userstory = prompt_user_stories | llm.with_structured_output(UserStories)
    response = chain_userstory.invoke({'requirements': state['requirements'], 'feedback_points': state.get('feedback_points', "")})
    state['user_stories'] = response.stories

    save_user_stories_to_txt(state, filename="user_stories.txt")
                             
    return state


def human_user_story_approval(state: State):
    # No operation – this is just a HITL pause node
    return state


def user_story_human_decision(state: State) -> Literal['Approve', 'Denied']:
    return state.get("user_story_status", "Approve")


def human_design_document_review(state: State):
    # No operation – this is just a HITL pause node
    return state


def design_document_human_decision(state:State) -> Literal['Approve', 'Denied']:
    return state.get("design_document_review_status", "Approve")


def save_design_document_to_word(state: dict, output_dir="artifacts", filename="design_document.docx"):
    """
    Saves the functional and technical design documents in a formatted Word (.docx) file.
    """
    if hasattr(state, "values") and not isinstance(state, dict):
        print("🔁 Detected .values object; extracting underlying dict")
        state = state.values

    design_doc = state.get("design_document", {})
    functional = design_doc.get("functional", [])
    technical = design_doc.get("technical", [])
    assumptions = design_doc.get("assumptions", [])
    open_questions = design_doc.get("open_questions", [])

    if not functional and not technical:
        print("⚠️ No design document data found in state.")
        return

    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    def add_section(doc, title, content):
        doc.add_heading(title, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item.strip(), style='List Bullet')
        elif isinstance(content, str):
            for line in content.strip().split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph("⚠️ Invalid content format")

    # Title
    doc.add_heading("Design Document", 0)

    # Sections
    add_section(doc, "Functional Requirements", functional)
    add_section(doc, "Technical Requirements", technical)
    add_section(doc, "Assumptions", assumptions)
    add_section(doc, "Open Questions / Risks", open_questions)

    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"✅ Design document saved to: {filepath}")


def create_design_document(state: State):

    prompt_create_design_document = PromptTemplate(
        template="""
        You are a senior software architect responsible for producing **detailed, production-grade design documents** based on a set of user stories, and optionally, prior review feedback.

        ### Inputs:
        1. **User Stories** (primary source of requirements)
        2. **Review Feedback** (optional, but if present, must be reflected in the output)

        ---

        ### Context:
        {user_stories}

        ### Feedback (if provided):
        {feedback}

        ---

        ### Task:
        Create a comprehensive **Design Document** that:
        - Clearly explains what the system should do from both user and technical perspectives
        - Incorporates review feedback wherever applicable
        - Adds depth, structure, and technical clarity suitable for developers, testers, and stakeholders

        ---

        ### Output Structure:

        #### 1. Functional Requirements
        - List and explain each major feature or capability from the end-user’s perspective.
        - Include:
            - Feature Name
            - Purpose
            - Inputs / Preconditions
            - Outputs / Postconditions
            - Acceptance Criteria
            - Edge Cases
            - User journey if applicable
        - Cover role-based behavior (e.g., Admin vs. Customer)
        - If feedback is given, update any feature logic or description based on it.

        #### 2. Technical Requirements
        - Detail system architecture:
            - Modules/components and their interactions
            - Technology stack and tools (backend, frontend, DB, etc.)
            - Deployment strategy (e.g., Docker, CI/CD, cloud infra)
        - API specifications:
            - Endpoint
            - HTTP Method
            - Request and Response JSON schemas
        - Database design:
            - Tables or collections
            - Relationships and example entries
        - Security:
            - Authentication and Authorization
            - Input validation, encryption, and access control
        - Performance and scalability
        - Add pseudocode or class structure where relevant
        - Improve clarity and modularity based on feedback if applicable.

        #### 3. Assumptions
        - List key assumptions made during design
        - E.g., "Payment service is already integrated", or "Admin has access to all bookings"

        #### 4. Open Questions / Risks
        - Identify any ambiguity, missing info, or dependencies
        - E.g., "Flight seat map integration pending", or "Query performance under load unknown"

        ---

        ### Output Format:
        Use clear section headers:

        **Functional Requirements:**
        <detailed feature bullets>

        **Technical Requirements:**
        <implementation details, APIs, data models>

        **Assumptions:**
        <list>

        **Open Questions / Risks:**
        <list>

        If no feedback is provided, ignore that section. If feedback exists, incorporate it meaningfully in both functional and technical sections.
        """,
        input_variables=["user_stories", "feedback"]
    )

    chain_create_design_document = prompt_create_design_document | llm.with_structured_output(DesignDocument)
    response = chain_create_design_document.invoke(
                            { "user_stories": "\n".join(state['user_stories']), 
                              "feedback": state.get("design_document_review_feedback", "")
                             })
    
    state['design_document'] = {
                                'functional': response.functional,
                                'technical': response.technical,
                                'assumptions': getattr(response, "assumptions", []),
                                'open_questions': getattr(response, "open_questions", [])
                                }
    
    save_design_document_to_word(state)
    
    return state


def save_files(file_blocks, output_dir="generated_code"):
    """
    Save parsed code blocks to files in specified directory.
    """
    os.makedirs(output_dir, exist_ok=True)
    
    for file in file_blocks:
        filename = file.get("filename", "unnamed.py")
        code = file.get("code", "")
        filepath = os.path.join(output_dir, filename)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(code)

        print(f"✅ Saved: {filename}")


def parse_files_from_response(response_text: str):
    """
    Parses code blocks from LLM output that includes:
    Filename: <file>
    Code:
    ```python
    <code>
    ```
    """
    pattern = r"Filename:\s*(?P<filename>[\w_]+\.py)\s*Code:\s*```(?:python)?\s*(?P<code>.*?)```"
    matches = re.finditer(pattern, response_text, re.DOTALL)

    files = []
    for match in matches:
        files.append({
            "filename": match.group("filename").strip(),
            "code": match.group("code").strip()
        })

    # Fallback in case of single unnamed code block
    if not files:
        fallback = re.search(r"```(?:python)?\s*(.*?)```", response_text, re.DOTALL)
        if fallback:
            files.append({
                "filename": "main.py",
                "code": fallback.group(1).strip()
            })

    return files


def parse_files_from_response(response_text: str):
    """
    Parses code blocks from LLM-generated response containing multiple Python files
    formatted with 'Filename:' and 'Code:' markers.

    Returns:
        List[Dict] - Each dict contains 'filename' and 'code' keys.
    """
    # Matches well-formed file blocks like:
    # Filename: user_interface.py
    # Code:
    # ```python
    # <code>
    # ```
    pattern = r"Filename:\s*(?P<filename>[\w_]+\.py)\s*Code:\s*```(?:python)?\s*(?P<code>.*?)```"
    matches = list(re.finditer(pattern, response_text, re.DOTALL))

    files = []
    if matches:
        for match in matches:
            files.append({
                "filename": match.group("filename").strip(),
                "code": match.group("code").strip()
            })
    else:
        # Fallback if only one large unnamed code block exists
        fallback_match = re.search(r"```(?:python)?\s*(.*?)```", response_text, re.DOTALL)
        if fallback_match:
            files.append({
                "filename": "main.py",  # default fallback
                "code": fallback_match.group(1).strip()
            })
        else:
            print("❌ No code blocks found in response.")
    
    return files


def generate_code(state: State):
    
    if state.get('code_review_status') == "Denied":
        prompt_regenerate_code = PromptTemplate(
            template = 
            """
            You are a senior software architect and software engineer responsible for building modular, production-grade systems.

            ---

            ### Context:
            - Below is the latest **Design Document** to implement:
            {design_document}

            - The following **Code Review Feedback** was provided:
            {feedback}

            - Below is the **previous version of the code**:
            {previous_code}

            ---


            Generate a complete implementation that addresses all feedback and follows best practices for security and code quality.

            ---

            ### Your Task:
            Regenerate the Python codebase from scratch using the updated design, **incorporating all feedback and improving upon the previous code** where applicable.

            ---
            ### Output Goal:
            Split the project requirements into multiple Python files, each with a specific responsibility, and generate clean, professional code for each requirement.

            ---

            ### Rules to Follow Strictly:
            1. For **each file**, include:
            - A `Filename:` line specifying the file name (in `snake_case.py`)
            - A `Code:` block with the actual Python code in a fenced markdown block like:
                ```
                Filename: user_interface.py
                Code:
                ```python
                <full python code>
                ```
            2. DO NOT skip the `Filename:` or `Code:` tags — even if there's only one file.
            3. Ensure code blocks contain complete imports and logic.
            4. Use **one file per logical component**, such as:
            - `api.py` for route handling
            - `models.py` for data models
            - `services.py` for business logic
            - `config.py` for environment setup
            - `main.py` or `app.py` as the entry point

            ---

            ### Coding Guidelines:
            - Follow the **Single Responsibility Principle**
            - Include comments for complex logic
            - Follow Python naming conventions (snake_case for files and functions)
            - Add docstrings for each function and class
            - Avoid unnecessary libraries
            - Include exception handling where needed

            ---

            ### Example Output Format(strictly follow this for every file):
            ---
            Filename: <file_name.py>
            Code:
            ```python
            <Full Python code for this file>

            Important Rules:
            DO NOT include any explanations, introductions, or summaries.
            DO NOT add any text outside the specified format.
            Each file must have its own Filename and Code block as shown.
            Maintain proper Python indentation and formatting.
            Assume the generated files will be saved separately in a project folder.
            """,
                input_variables=["design_document", "feedback", "previous_code"]
            )

        chain_code_regeneration = prompt_regenerate_code | llm.with_structured_output(GenerateCode)
        code_response = chain_code_regeneration.invoke(
            {"design_document":state['design_document'],
             "feedback":state['code_review_feedback'],
             "previous_code":state['code'],
             })
    elif state.get('security_review_status') == "Denied" and state.get('code_review_status') == "Approve":
        prompt_regenerate_code = PromptTemplate(
            template = 
            """
            You are a senior software engineer and cybersecurity expert responsible for producing clean, modular, and secure production-grade Python code.

            ---

            ### Context:
            - Below is the latest **Design Document** to implement:
            {design_document}

            - The following **Security Review Feedback** was provided:
            {security_feedback}

            - Below is the **previous version of the code**:
            {previous_code}

            ---

            ### Your Objective:

            Regenerate the full Python codebase by:
            - Fully implementing the design specification
            - Fixing all security concerns mentioned in the feedback
            - Improving or reworking vulnerable logic in the old code
            - Ensuring every security best practice is followed
            ---
            ### Security Must-Haves:

            - Sanitize and validate all user inputs
            - Use secure authentication and session management practices
            - Avoid hardcoded secrets or credentials
            - Implement proper access controls and role checks
            - Prevent SQL injection, XSS, and command injection vulnerabilities
            - Include encryption for sensitive data where appropriate
            - Handle errors securely without leaking internal state

            ---
            ### Output Format (Strict — One Block Per File):

            ### Rules to Follow Strictly:
            1. For **each file**, include:
            - A `Filename:` line specifying the file name (in `snake_case.py`)
            - A `Code:` block with the actual Python code in a fenced markdown block like:
                ```
                Filename: user_interface.py
                Code:
                ```python
                <full python code>
                ```
            2. DO NOT skip the `Filename:` or `Code:` tags — even if there's only one file.
            3. Ensure code blocks contain complete imports and logic.
            4. Use **one file per logical component**, such as:
            - `api.py` for route handling
            - `models.py` for data models
            - `services.py` for business logic
            - `config.py` for environment setup
            - `main.py` or `app.py` as the entry point

            ---

            ### Coding Guidelines:
            - Follow the **Single Responsibility Principle**
            - Include comments for complex logic
            - Follow Python naming conventions (snake_case for files and functions)
            - Add docstrings for each function and class
            - Avoid unnecessary libraries
            - Include exception handling where needed

            ---

            ### Example Output Format(strictly follow this for every file):
            ---
            Filename: <file_name.py>
            Code:
            ```python
            <Full Python code for this file>

            Important Rules:
            DO NOT include any explanations, introductions, or summaries.
            DO NOT add any text outside the specified format.
            Each file must have its own Filename and Code block as shown.
            Maintain proper Python indentation and formatting.
            Assume the generated files will be saved separately in a project folder.
            """,
                input_variables=["design_document", "feedback", "previous_code"]
            )

        chain_code_regeneration = prompt_regenerate_code | llm.with_structured_output(GenerateCode)
        code_response = chain_code_regeneration.invoke(
            {"design_document":state['design_document'],
             "feedback":state['security_review_feedback'],
             "previous_code":state['code'],
             })
    elif state.get('security_review_status') == "Approve" and state.get('code_review_status') == "Approve" and state.get('qa_review_status') == "Denied":
        prompt_regenerate_code = PromptTemplate(
            template = 
            """
            You are a senior software engineer, expert quality professional and and cybersecurity expert responsible for producing clean, modular, and secure production-grade Python code.

            ---

            ### Context:
            - Below is the latest **Design Document** to implement:
            {design_document}

            - The following **QA Review Feedback** was provided:
            {qa_feedback}

            - The following **Security Review Feedback** was provided:
            {security_feedback}

            - Below is the **previous version of the code**:
            {previous_code}

            ---

            ### Your Objective:

            Regenerate the full Python codebase by:
            - Fully implementing the design specification
            - Fixing all security concerns mentioned in the feedback
            - Improving or reworking vulnerable logic in the old code
            - Ensuring every security best practice is followed
            ---
            ### Security Must-Haves:

            - Sanitize and validate all user inputs
            - Use secure authentication and session management practices
            - Avoid hardcoded secrets or credentials
            - Implement proper access controls and role checks
            - Prevent SQL injection, XSS, and command injection vulnerabilities
            - Include encryption for sensitive data where appropriate
            - Handle errors securely without leaking internal state

            ---
            ### Output Format (Strict — One Block Per File):

            ### Rules to Follow Strictly:
            1. For **each file**, include:
            - A `Filename:` line specifying the file name (in `snake_case.py`)
            - A `Code:` block with the actual Python code in a fenced markdown block like:
                ```
                Filename: user_interface.py
                Code:
                ```python
                <full python code>
                ```
            2. DO NOT skip the `Filename:` or `Code:` tags — even if there's only one file.
            3. Ensure code blocks contain complete imports and logic.
            4. Use **one file per logical component**, such as:
            - `api.py` for route handling
            - `models.py` for data models
            - `services.py` for business logic
            - `config.py` for environment setup
            - `main.py` or `app.py` as the entry point

            ---

            ### Coding Guidelines:
            - Follow the **Single Responsibility Principle**
            - Include comments for complex logic
            - Follow Python naming conventions (snake_case for files and functions)
            - Add docstrings for each function and class
            - Avoid unnecessary libraries
            - Include exception handling where needed

            ---

            ### Example Output Format(strictly follow this for every file):
            ---
            Filename: <file_name.py>
            Code:
            ```python
            <Full Python code for this file>

            Important Rules:
            DO NOT include any explanations, introductions, or summaries.
            DO NOT add any text outside the specified format.
            Each file must have its own Filename and Code block as shown.
            Maintain proper Python indentation and formatting.
            Assume the generated files will be saved separately in a project folder.
            """,
                input_variables=["design_document", "qa_feedback", "security_feedback", "previous_code"]
            )

        chain_code_regeneration = prompt_regenerate_code | llm
        # Prepare input with proper handling of optional fields
        input_data = {
            "design_document": state['design_document'],
            "qa_feedback": '\n'.join(state.get('qa_review_feedback', [])) if isinstance(state.get('qa_review_feedback'), list) else str(state.get('qa_review_feedback', '')),
            "security_feedback": '\n'.join(state.get('security_review_feedback', [])) if isinstance(state.get('security_review_feedback'), list) else str(state.get('security_review_feedback', '')),
            "previous_code": str(state.get('code', ''))
        }
        code_response = chain_code_regeneration.invoke(input_data)
        code_response = code_response.content if hasattr(code_response, 'content') else code_response
        
    else:
        prompt_generate_code = PromptTemplate(
            template = 
            """
            You are a senior software architect and software engineer responsible for building modular, production-grade systems.

            Your task is to generate Python code based **only** on the following design document:

            {design_document}

            ---

            ### Output Goal:
            Split the project requirements into multiple Python files, each with a specific responsibility, and generate clean, professional code for each requirement.

            ---

            ### Rules to Follow Strictly:
            1. For **each file**, include:
            - A `Filename:` line specifying the file name (in `snake_case.py`)
            - A `Code:` block with the actual Python code in a fenced markdown block like:
                ```
                Filename: user_interface.py
                Code:
                ```python
                <full python code>
                ```
            2. DO NOT skip the `Filename:` or `Code:` tags — even if there's only one file.
            3. Ensure code blocks contain complete imports and logic.
            4. Use **one file per logical component**, such as:
            - `api.py` for route handling
            - `models.py` for data models
            - `services.py` for business logic
            - `config.py` for environment setup
            - `main.py` or `app.py` as the entry point

            ---

            ### Coding Guidelines:
            - Follow the **Single Responsibility Principle**
            - Include comments for complex logic
            - Follow Python naming conventions (snake_case for files and functions)
            - Add docstrings for each function and class
            - Avoid unnecessary libraries
            - Include exception handling where needed

            ---

            ### Example Output Format(strictly follow this for every file):
            ---
            Filename: <file_name.py>
            Code:
            ```python
            <Full Python code for this file>

            Important Rules:
            DO NOT include any explanations, introductions, or summaries.
            DO NOT add any text outside the specified format.
            Each file must have its own Filename and Code block as shown.
            Maintain proper Python indentation and formatting.
            Assume the generated files will be saved separately in a project folder.
            """,
                input_variables={"design_document"}
            )
        # chain_code_generation = prompt_generate_code | llm.with_structured_output(GenerateCode)
        # code_response = chain_code_generation.invoke({"design_document":state['design_document']})
        chain_code_generation = prompt_generate_code | llm
        code_response = chain_code_generation.invoke({"design_document": state['design_document']})
    
    generated_code = code_response.content if hasattr(code_response, "content") else code_response


    print("==> code_response: ", generated_code)

    try:
        # Parse the code blocks from the response
        file_blocks = parse_files_from_response(generated_code)

        # Save the files
        save_files(file_blocks)

        # Save the generated code to state
        state['code'] = generated_code
    except Exception as e:
        print(f"Error parsing code response: {e}")
        # If parsing fails, save the raw response
        state['code'] = str(generated_code)

    return state


def human_code_review(state: State):
    # No operation – this is just a HITL pause node
    return state


def code_review_human_decision(state:State) -> Literal['Approve', 'Denied']:
    return state.get("code_review_status", "Approve")


def security_review(state: State):
    """Conducts a security review of the code to check for vulnerabilities."""
    prompt_security = PromptTemplate(
        template="""You are a senior cybersecurity expert specializing in secure coding practices and vulnerability assessment.

        Task: Conduct a thorough security review of the following code:

        **Code:**
        {generated_code}

        Provide structured feedback, including detected issues and suggested fixes.
        Format:
        - Status: Approve / Denied
        - Feedback: (Explain security risks and provide recommended changes)
        
        """,
        input_variables=["generated_code"]
        )
    
    chain_security = prompt_security | llm.with_structured_output(Review)
    response_security = chain_security.invoke({
        "generated_code": state['code']
    })

    state['security_review_status'] = response_security.status
    state['security_review_feedback'] = response_security.review

    return state

def human_security_review(state: State):
    # No operation – this is just a HITL pause node
    return state


def security_review_human_decision(state:State) -> Literal['Approve', 'Denied']:
    return state.get("security_review_status", "Approve")


def human_test_cases_review(state: State):
    # No operation – this is just a HITL pause node
    return state


def test_cases_review_human_decision(state:State) -> Literal['Approve', 'Denied']:
    return state.get("test_cases_review_status", "Approve")



def save_test_cases_to_files(response_text, output_dir="test_cases"):
    os.makedirs(output_dir, exist_ok=True)

    # Extract string content from AIMessage if necessary
    if hasattr(response_text, "content"):
        response_text = response_text.content

    test_cases = re.split(r"\n---\n", response_text.strip())

    for case in test_cases:
        name_match = re.search(r"\[Test Case Name\]: (.+)", case)
        if name_match:
            name = name_match.group(1).strip()
            safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', name)
            filepath = os.path.join(output_dir, f"{safe_name}.txt")
            with open(filepath, "w") as f:
                f.write(case.strip())
            print(f"✅ Saved: {filepath}")


def write_test_cases(state:State):
    """Generates test cases for the code based on functional and technical design documents."""
    
    # Validate required state
    if "code" not in state or not state["code"]:
        raise KeyError("❌ 'code' not found in state or is empty. Code must be generated before writing test cases.")
    if "design_document" not in state or not state["design_document"]:
        raise KeyError("❌ 'design_document' not found in state or is empty. Design document must be created before writing test cases.")
    
    if state.get("test_cases_review_status") == "Denied":
        prompt_test_case_rewrite = PromptTemplate(
            template="""You are a senior QA engineer responsible for writing high-quality test cases for Python systems.

            ---

            ### Objective:
            Based on the **previous text cases**, **feedback**, **design specifications** and **generated code**, rewrite a comprehensive suite of test cases to cover all the points from feedback and refactor the code for important functionalities, edge cases, and failure paths as per the feedback.

            ---
            ### Inputs:
            **Previous Text Cases:**
            {old_test_cases}

            **Feedback:**
            {test_cases_review_feedback}

            **Generated Code:**
            {generated_code}

            **Design Document:**
            {design_document}
            ---
            ### Requirements:

            1. Analyze each functional and technical requirement.
            2. For **each major functionality**, generate:
            - One happy path (positive) test case
            - One edge case test
            - One negative/failure test
            3. Cover:
            - Validations
            - Business rules
            - Exception handling
            - Security conditions if applicable

            ---

            ### Format for Each Test Case (Strict):

            Each test case should be formatted exactly like this, **repeated multiple times in the same response**:

            [Test Steps]:

            1. Step 1

            2. Step 2
            ---

            ### Output Rules:
            - Output as a **single continuous string** with multiple test case blocks.
            - Separate each test case with a line: `---`
            - DO NOT wrap it in Markdown or JSON.
            - DO NOT include summaries, explanations, or file names — only test case blocks.
            
            ---
            Example Output:
            [Test Case Name]: Search flights - valid input

            [Description]:
            Tests whether flight search returns correct results with valid input.

            [Test Type]: Unit

            [Test Steps]:
            1. Send GET request to /flights?destination=NYC&date=2024-12-01
            2. Mock database with sample flights
            3. Receive response

            [Expected Result]:
            List of flights for NYC on specified date

            ---
            [Test Case Name]: Search flights - invalid date format

            [Description]:
            Validates system's behavior when date is malformed.

            [Test Type]: Negative

            [Test Steps]:
            1. Send GET request to /flights?date=12-01-2024
            2. Observe API response

            [Expected Result]:
            400 Bad Request with validation error messag
            """,
            input_variables=["generated_code", "test_cases_review_feedback", "design_document", "old_test_cases"]
        )

        chain_test_case_rewrite = prompt_test_case_rewrite | llm
        test_cases = chain_test_case_rewrite.invoke({
            "generated_code": state["code"],
            "design_document": state["design_document"],
            "old_test_cases": state["test_cases"],
            "test_cases_review_feedback": state["test_cases_review_feedback"],
            
        })
    else:
        prompt_test_case = PromptTemplate(
            template="""You are a senior QA engineer responsible for writing high-quality test cases for Python systems.

            ---

            ### Objective:
            Based on the **design specifications** and **generated code**, write a comprehensive suite of test cases to cover all important functionalities, edge cases, and failure paths.

            ---
            ### Inputs:
            **Generated Code:**
            {generated_code}

            **Design Document:**
            {design_document}
            ---
            ### Requirements:

            1. Analyze each functional and technical requirement.
            2. For **each major functionality**, generate:
            - One happy path (positive) test case
            - One edge case test
            - One negative/failure test
            3. Cover:
            - Validations
            - Business rules
            - Exception handling
            - Security conditions if applicable

            ---

            ### Format for Each Test Case (Strict):

            Each test case should be formatted exactly like this, **repeated multiple times in the same response**:

            [Test Steps]:

            1. Step 1

            2. Step 2
            ---

            ### Output Rules:
            - Output as a **single continuous string** with multiple test case blocks.
            - Separate each test case with a line: `---`
            - DO NOT wrap it in Markdown or JSON.
            - DO NOT include summaries, explanations, or file names — only test case blocks.
            
            ---
            Example Output:
            [Test Case Name]: Search flights - valid input

            [Description]:
            Tests whether flight search returns correct results with valid input.

            [Test Type]: Unit

            [Test Steps]:
            1. Send GET request to /flights?destination=NYC&date=2024-12-01
            2. Mock database with sample flights
            3. Receive response

            [Expected Result]:
            List of flights for NYC on specified date

            ---
            [Test Case Name]: Search flights - invalid date format

            [Description]:
            Validates system's behavior when date is malformed.

            [Test Type]: Negative

            [Test Steps]:
            1. Send GET request to /flights?date=12-01-2024
            2. Observe API response

            [Expected Result]:
            400 Bad Request with validation error messag
            """,
            input_variables=["generated_code", "design_document"]
        )

        chain_test_case = prompt_test_case | llm
        test_cases = chain_test_case.invoke({
            "generated_code": state["code"],
            "design_document": state["design_document"]
        })

    # Extract test cases content if it's an AIMessage
    raw_test_cases = getattr(test_cases, "content", test_cases)
    if isinstance(raw_test_cases, dict) and "text" in raw_test_cases:
        raw_test_cases = raw_test_cases["text"]

    # Validate test cases type
    if not isinstance(raw_test_cases, str):
        raise TypeError(f"Expected test_cases to be str or AIMessage with .content, but got: {type(raw_test_cases)}")

    # Save test cases to files
    save_test_cases_to_files(raw_test_cases, output_dir="test_cases")

    # Store in state
    state["test_cases"] = raw_test_cases
    print("✅ Test cases stored in state:", state["test_cases"][:100], "...")
    
    return state


def human_qa_review(state: State):
    # No operation – this is just a HITL pause node
    return state


def qa_review_human_decision(state:State) -> Literal['Approve', 'Denied']:
    if state.get('qa_review_status') == 'Approve':
        return 'Approve'
    return 'Denied'


def qa_testing(state: State):
    """Conducts QA testing."""
    print("==> State", state)

    # Validate required state
    if "code" not in state or not state["code"]:
        raise KeyError("❌ 'code' not found in state or is empty. Code must be generated before QA Testing.")
    if "test_cases" not in state or not state["test_cases"]:
        raise KeyError("❌ 'test_cases' not found in state or is empty. Make sure test cases are generated before QA Testing.")

    # Extract content safely
    code = getattr(state["code"], "content", state["code"]) if isinstance(state["code"], object) else state["code"]
    testcases = getattr(state["test_cases"], "content", state["test_cases"]) if isinstance(state["test_cases"], object) else state["test_cases"]

    prompt_qa_test = PromptTemplate(
        template="""You are a seasoned QA engineer with expertise in thorough testing and quality validation.

        Task: Perform a comprehensive QA evaluation of the following system.

        ---
        ### Code:
        ```python
        {code}

        Test Cases:
        {testcases} 
        
        ---
        
        Evaluate whether the test cases adequately cover the code.

        Simulate test execution and report the outcome for each test (if possible).

        Return a final status as either: Approve or Denied.

        Include concise feedback under review with specific suggestions or validation results.
        """,
        input_variables=["code","testcases"]
    )

    chain_qa_test = prompt_qa_test | llm.with_structured_output(Review)
    response = chain_qa_test.invoke({"code":code, "testcases":testcases})
    state["qa_review_status"] = response.status
    state["qa_review_feedback"] = response.review
    
    return state


def deployment(state: State):
    if state.get('qa_review_status') == 'Approve':
        state['deployment'] = 'deployed'
    return state

# Design a graph
from langgraph.graph import END, StateGraph, START


graph_builder = StateGraph(State)

# Define the nodes
graph_builder.add_node("User Requirements", user_input_requirements)
graph_builder.add_node("Auto-generate User Stories", auto_generate_user_stories)
graph_builder.add_node("Human User Story Approval", human_user_story_approval)
graph_builder.add_node("Create Design Document", create_design_document)
graph_builder.add_node("Human Design Document Review", human_design_document_review)
graph_builder.add_node("Generate Code", generate_code)
graph_builder.add_node("Human Code Review", human_code_review)
graph_builder.add_node("Security Review", security_review)
graph_builder.add_node("Human Security Review", human_security_review)
graph_builder.add_node("Write Test Cases", write_test_cases)
graph_builder.add_node("Human Test Cases Review", human_test_cases_review)
graph_builder.add_node("QA Testing", qa_testing)
graph_builder.add_node("Human QA Review", human_qa_review)
graph_builder.add_node("Deployment", deployment)

# graph_builder.add_node("Fix Code after QA Feedback", fix_code_after_qa_feedback)
# # graph_builder.add_node("Monitoring", monitoring)
# # graph_builder.add_node("Requirement Change", requirement_change)




graph_builder.add_edge(START, "User Requirements")
graph_builder.add_edge("User Requirements", "Auto-generate User Stories")
graph_builder.add_edge("Auto-generate User Stories", "Human User Story Approval")
graph_builder.add_conditional_edges(
    "Human User Story Approval",
    user_story_human_decision,
    {
        "Approve": "Create Design Document",
        "Denied": "Auto-generate User Stories"
    }
)
graph_builder.add_edge("Create Design Document", "Human Design Document Review")
graph_builder.add_conditional_edges(
    "Human Design Document Review",
    design_document_human_decision,
    {
        "Approve": "Generate Code",
        "Denied": "Create Design Document"
    }
)
graph_builder.add_edge("Generate Code", "Human Code Review")
graph_builder.add_conditional_edges(
    "Human Code Review",
    code_review_human_decision,
    {
        "Approve": "Security Review",
        "Denied": "Generate Code"
    }
)
graph_builder.add_edge("Security Review", "Human Security Review")
graph_builder.add_conditional_edges(
    "Human Security Review",
    security_review_human_decision,
    {
        "Approve": "Write Test Cases",
        "Denied": "Generate Code"
    }
)
graph_builder.add_edge("Write Test Cases", "Human Test Cases Review")
graph_builder.add_conditional_edges(
    "Human Test Cases Review",
    test_cases_review_human_decision,
    {
        "Approve": "QA Testing",
        "Denied": "Write Test Cases"
    }
)
graph_builder.add_edge("QA Testing", "Human QA Review")
graph_builder.add_conditional_edges(
    "Human QA Review",
    qa_review_human_decision,
    {
        "Approve": "Deployment",
        "Denied": "Generate Code"
    }
)
graph_builder.add_edge("Deployment", END)


# compile the graph
memory = MemorySaver()
graph = graph_builder.compile(interrupt_before=["Human User Story Approval", "Human Design Document Review", "Human Code Review", "Human Security Review", "Human Test Cases Review", "Human QA Review"], checkpointer=memory)


# Save the PNG to a file
# with open('react_graph.mmd', 'wb') as f:
#     f.write(graph.get_graph().draw_mermaid_png())
# print("Graph image saved as 'react_graph.png'. Open it to view the graph.")
with open('react_graph.mmd', 'w') as f:
    f.write(graph.get_graph().draw_mermaid())
print("Graph Mermaid syntax saved as 'react_graph.mmd'. You can paste this into https://mermaid.live to view the graph.")



thread = {
    "configurable":{
        "thread_id": str(uuid.uuid4())
    }
}

__all__ = ["State", "graph"]
